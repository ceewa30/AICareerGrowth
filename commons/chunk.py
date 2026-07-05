import os
import docx
from typing import List
# Assuming your definitions match your local data_validation import environment
import sys
parent_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)
from core.VectorEmbedding import VectorEmbedding
from commons.DataValidation import UserProfile, RepoDocument, READMEChunk, ResumeExtractionSchema
import glob
from pathlib import Path
import docx
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv(override=True)


def extract_raw_resume_data(resume_text: str) -> ResumeExtractionSchema:
    """Uses OpenAI to extract raw unstructured text data into a structured schema."""
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    completion = client.beta.chat.completions.parse(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Extract core career metrics from the resume text parameters carefully."},
            {"role": "user", "content": f"Parse the following resume context:\n{resume_text}"}
        ],
        response_format=ResumeExtractionSchema,
    )
    return completion.choices[0].message.parsed


def recursive_text_splitter(text: str, chunk_size: int = 800, chunk_overlap: int = 150) -> List[str]:
    """
    Pure text utility that recursively splits text into chunks below the specified size limit.
    """
    print(f"DEBUG: Processing block length: {len(text)}")

    if chunk_overlap >= chunk_size:
        chunk_overlap = chunk_size // 2

    text = text.strip()

    if len(text) <= chunk_size:
        return [text]

    separators = ["\n\n", "\n", " "]
    chosen_separator = None
    for sep in separators:
        if sep in text:
            chosen_separator = sep
            break

    if chosen_separator is None:
        chunks = []
        step = chunk_size - chunk_overlap
        for i in range(0, len(text), step):
            slice_chunk = text[i:i + chunk_size]
            if slice_chunk:
                chunks.append(slice_chunk)
        return chunks

    parts = text.split(chosen_separator)
    final_chunks = []
    current_chunk = []
    current_length = 0

    for part in parts:
        part_clean = part.strip()
        if not part_clean:
            continue

        if len(part_clean) > chunk_size:
            if current_chunk:
                final_chunks.append(chosen_separator.join(current_chunk))
                current_chunk = []
                current_length = 0

            step = chunk_size - chunk_overlap
            for i in range(0, len(part_clean), step):
                slice_segment = part_clean[i:i + chunk_size]
                if slice_segment:
                    final_chunks.append(slice_segment)
            continue

        part_len = len(part) + (len(chosen_separator) if current_chunk else 0)
        if current_length + part_len <= chunk_size:
            current_chunk.append(part)
            current_length += part_len
        else:
            if current_chunk:
                final_chunks.append(chosen_separator.join(current_chunk))

            # Compute overlap buffer safely
            overlap_buffer = []
            overlap_len = 0
            for p in reversed(current_chunk):
                p_len = len(p) + (len(chosen_separator) if overlap_buffer else 0)
                if overlap_len + p_len <= chunk_overlap:
                    overlap_buffer.insert(0, p)
                    overlap_len += p_len
                else:
                    break
            current_chunk = overlap_buffer + [part]
            current_length = sum(len(p) for p in current_chunk) + (len(chosen_separator) * (len(current_chunk) - 1))

    if current_chunk:
        final_chunks.append(chosen_separator.join(current_chunk))

    # Clean up oversized fragments recursively using the same inner text parameters
    verified_chunks = []
    for chunk in final_chunks:
        if len(chunk) > chunk_size:
            # Safe to run a shallow fallback split since we know individual words are already broken down
            step = chunk_size - chunk_overlap
            for i in range(0, len(chunk), step):
                verified_chunks.append(chunk[i:i + chunk_size])
        else:
            verified_chunks.append(chunk)

    return [c.strip() for c in verified_chunks if c.strip()]

if __name__ == "__main__":
    embedding_client = VectorEmbedding()

    knowledge = {}
    filenames = glob.glob("../resources/*.docx")
    print(f"Glob pulling : {filenames}")
    try:
        for filename in filenames:
            name = Path(filename).stem.split(' ')[-1]

            doc = docx.Document(filename)
            file_text = "\n".join([p.text for p in doc.paragraphs if p.text])

            # Store the text using the lowercased name as the key
            knowledge[name.lower()] = file_text

        print(f"Loaded knowledge keys: {list(knowledge.keys())}")

        structured_readme_chunks = []
        for doc_name, doc_text in knowledge.items():
            print(f"\n ----- Chunking Resource: {doc_name} -----")
            raw_resume_data = extract_raw_resume_data(doc_text)

            chunks = recursive_text_splitter(doc_text, chunk_size=800, chunk_overlap=150)
            print(f"✅ Generated {len(chunks)} chunks for '{doc_name}'")

            vectors = embedding_client.generate_batch_embeddings(chunks)
            print(f"🧬 Generated {vectors} matching 1536-dim vector arrays successfully.")

            for idx, chunk_text in enumerate(chunks):
                chunk_metadata = {
                    "source_doc": doc_name,
                    "chunk_index": idx,
                    "file_type": "docx"
                }
                chunk_obj = READMEChunk(
                    chunk_id=f"{doc_name}_chunk_{idx}", # Unique tracking identifier
                    content=chunk_text,                 # The actual chunk text string
                    metadata=chunk_metadata,
                    chunk_index=idx,
                    embedding=vectors[idx]
                )
                structured_readme_chunks.append(chunk_obj)

                print(f"Chunking is {chunk_obj}")

                repo_doc_asset = RepoDocument(
                    repo_name=f"{doc_name}_resume_source",
                    raw_readme=doc_text,
                    readme_chunks=structured_readme_chunks
                )

                final_user_profile = UserProfile(
                    name=raw_resume_data.name,
                    current_role=raw_resume_data.current_role,
                    github_skills=set(raw_resume_data.extracted_skills),
                    linkedin_skills=set(),
                    linkedin_summary=raw_resume_data.professional_summary
                )

            print(f"\n🏆 Successfully Created Global Profile for: {final_user_profile.name}")
            print(f"💼 Extracted Current Role Title: {final_user_profile.current_role}")

            # Safe Fallback: Check model dictionary keys if property access drops out
            dumped_profile = final_user_profile.model_dump()

            # Read from serialized output or fallback to manual generation
            chunks_count = dumped_profile.get("total_chunks_count")
            skills_count = dumped_profile.get("all_skills", final_user_profile.all_skills)

            print(f"📊 Computed Property Verification -> Total Chunks Tracked: {chunks_count}")
            print(f"💡 Normalized Unified Skills Set Count: {skills_count}")
            print(f"🔍 Normalized Skills Array Preview: {list(final_user_profile.all_skills)[:5]}")

        print(f"\n📊 Validation Layer Complete: Successfully created {len(structured_readme_chunks)} READMEChunk objects.")

    except FileNotFoundError:
        print(f"❌ Error: The file '{filenames}' was not found.")
    except Exception as e:
        print(f"❌ Ingestion Error: Failed to parse uploaded file. Detail: {str(e)}")
