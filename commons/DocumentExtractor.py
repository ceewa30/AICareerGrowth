import os
import sys
from pathlib import Path
import docx
import glob
import re
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from core.VectorEmbedding import VectorEmbedding
from commons.DataValidation import UserProfile, READMEChunk, RepoDocument, ResumeExtractionSchema, DynamicRepoMetadata
from commons.chunk import recursive_text_splitter
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv(override=True)

def extract_raw_resume_data(resume_text: str) -> DynamicRepoMetadata:
    """Uses OpenAI to extract raw unstructured text data into a structured schema."""
    root_dir = Path(__file__).resolve().parent
    if (root_dir / "resources").exists():
        project_root = root_dir
    else:
        project_root = root_dir.parent

    prompt_path = project_root / "resources" / "system_prompt.txt"
    print(f"Loading prompt from: {prompt_path}")

    # Open and read the raw file content
    with open(prompt_path, "r", encoding="utf-8") as f:
        raw_content = f.read().strip()

    pattern = r'"document"\s*:\s*"(.*?)"\s*,\s*"github"'
    match = re.search(pattern, raw_content, re.DOTALL)

    if match:
        system_prompt = match.group(1).strip()
        # print("\n--- Extracted Document Prompt Safely ---")
        # print(system_prompt)
    else:
        print("Error: Could not locate the 'document' prompt structure in your file.")

    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    completion = client.beta.chat.completions.parse(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Parse the following resume context:\n{resume_text}"}
        ],
        response_format=DynamicRepoMetadata,
    )
    return completion.choices[0].message.parsed


class DocumentExtractor:
    def __init__(self):
        self.root_dir = Path(__file__).resolve().parent
        if (self.root_dir / "resources").exists():
            self.project_root = self.root_dir
        else:
            self.project_root = self.root_dir.parent

        self.embedding_client = VectorEmbedding()

        search_path = str(self.project_root / "resources" / "*.docx")
        self.filenames = glob.glob(search_path)
        print(f"Glob pulling : {self.filenames}")

    def process_file_upload(self) -> str:
        """Helper to detect file types and parse text out of incoming Gradio file objects safely."""

        knowledge = {}
        documents = []
        metadatas = []
        ids = []
        embeddings = []

        documents.clear()
        metadatas.clear()
        ids.clear()
        embeddings.clear()

        # Track the global collection of all structured document model structures
        all_repo_documents = []
        try:
            for filename in self.filenames:
                name = Path(filename).stem.split(' ')[-1]

                doc = docx.Document(filename)
                file_text = "\n".join([p.text for p in doc.paragraphs if p.text])

                # Store the text using the lowercased name as the key
                knowledge[name.lower()] = file_text
                print(f"Loaded knowledge keys: {list(knowledge.keys())}")

            structured_readme_chunks = []
            for doc_name, doc_text in knowledge.items():
                print(f"\n ----- Chunking Resource: {doc_name} -----")
                raw_resume_data = extract_raw_resume_data(doc_text)

                flat_profile_metadata = {
                    "repository_name": "docx_resume_source",
                    "author": raw_resume_data.author,
                    "primary_stack": raw_resume_data.tech_stack,
                    "ml_used": raw_resume_data.uses_machine_learning,
                    "algo_type": raw_resume_data.algorithm_summary,
                    "current_role": raw_resume_data.current_role,
                    "file_type": "docx"
                }
                print(f"Metadata : {flat_profile_metadata}")

                chunks = recursive_text_splitter(doc_text, chunk_size=800, chunk_overlap=150)
                print(f"✅ Generated {len(chunks)} chunks for '{doc_name}'")

                chunk_vectors = self.embedding_client.generate_batch_embeddings(chunks)
                print(f"🧬 Generated {len(chunk_vectors)} matching 1536-dim vector arrays successfully.")

                current_doc_chunks = []

                for idx, chunk_text in enumerate(chunks):

                    row_id = f"docx_chunk_{idx}".replace('/', '_').replace('-', '_')

                    chunk_specific_meta = flat_profile_metadata.copy()
                    chunk_specific_meta["chunk_index"] = idx

                    chunk_obj = READMEChunk(
                        chunk_id=f"{doc_name}_chunk_{idx}",
                        content=chunk_text,
                        metadata=chunk_specific_meta,
                        chunk_index=idx,
                        embedding=chunk_vectors[idx]
                    )
                    current_doc_chunks.append(chunk_obj)
                    structured_readme_chunks.append(chunk_obj)

                    documents.append(chunk_text)
                    metadatas.append(chunk_specific_meta)
                    ids.append(row_id)
                    embeddings.append(chunk_vectors[idx])

                    repo_doc_asset = RepoDocument(
                        repo_name=f"{doc_name}_resume_source",
                        raw_readme=doc_text,
                        readme_chunks=current_doc_chunks
                    )
                    all_repo_documents.append(repo_doc_asset)

            docx_payload = {
                "documents": documents,
                "metadatas": metadatas,
                "ids": ids,
                "embeddings": embeddings,
                "repo_documents": all_repo_documents
            }

            return docx_payload

        except FileNotFoundError:
            print(f"❌ Error: The file '{self.filenames}' was not found.")
        except Exception as e:
            print(f"❌ Ingestion Error: Failed to parse uploaded file. Detail: {str(e)}")



if __name__=="__main__":
    extractor = DocumentExtractor()

    # Run parsing extraction test checks against your local doc path instance
    extracted_text = extractor.process_file_upload()
